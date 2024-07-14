import heapq
from collections import Counter
import graphviz
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Node:
    def __init__(self, char, freq):
        self.char = char
        self.freq = freq
        self.left = None
        self.right = None
    
    def __lt__(self, other):
        return self.freq < other.freq

def create_huffman_tree(text):
    freq = Counter(text)
    heap = [Node(char, count) for char, count in freq.items()]
    heapq.heapify(heap)
    steps = []
    while len(heap) > 1:
        left = heapq.heappop(heap)
        right = heapq.heappop(heap)
        merged = Node(None, left.freq + right.freq)
        merged.left = left
        merged.right = right
        heapq.heappush(heap, merged)
        steps.append((left, right, merged))
    return heap[0], steps

def generate_codes(node, code='', mapping={}):
    if node.char:
        mapping[node.char] = code
    else:
        generate_codes(node.left, code + '0', mapping)
        generate_codes(node.right, code + '1', mapping)
    return mapping

def visualize_tree(node, dot, node_id):
    if node.char:
        dot.node(str(node_id), f'{node.char}\n{node.freq}')
    else:
        dot.node(str(node_id), f'{node.freq}')
        left_id = node_id * 2
        right_id = node_id * 2 + 1
        dot.edge(str(node_id), str(left_id), '0')
        dot.edge(str(node_id), str(right_id), '1')
        visualize_tree(node.left, dot, left_id)
        visualize_tree(node.right, dot, right_id)

def create_step_image(step, index):
    dot = graphviz.Digraph()
    dot.attr(rankdir='TB', size='3,3')
    left, right, merged = step
    dot.node('left', f'{left.char if left.char else ""}\n{left.freq}')
    dot.node('right', f'{right.char if right.char else ""}\n{right.freq}')
    dot.node('merged', f'{merged.freq}')
    dot.edge('merged', 'left', '0')
    dot.edge('merged', 'right', '1')
    filename = f'langkah_{index}'
    dot.render(filename, format='png', cleanup=True)
    return filename + '.png'

def create_final_tree_image(root):
    dot = graphviz.Digraph()
    dot.attr(rankdir='TB', size='4,4')
    visualize_tree(root, dot, 1)
    filename = 'pohon_akhir'
    dot.render(filename, format='png', cleanup=True)
    return filename + '.png'

def calculate_compression(text, codes):
    original_size = len(text) * 8
    compressed_size = sum(len(codes[char]) for char in text)
    compression_ratio = compressed_size / original_size
    space_saving = 1 - compression_ratio
    return compression_ratio, space_saving

def create_word_document(text, steps, root, codes):
    doc = Document()
    doc.add_heading('Proses Pembuatan Kode Huffman', 0)
    
    doc.add_paragraph(f"Teks asli: '{text}'")
    doc.add_paragraph("Frekuensi kemunculan setiap karakter:")
    freq_table = doc.add_table(rows=1, cols=3)
    freq_table.style = 'Table Grid'
    hdr_cells = freq_table.rows[0].cells
    hdr_cells[0].text = 'Karakter'
    hdr_cells[1].text = 'Frekuensi'
    hdr_cells[2].text = 'Penjelasan'
    for char, count in Counter(text).items():
        row_cells = freq_table.add_row().cells
        row_cells[0].text = f"'{char}'"
        row_cells[1].text = str(count)
        row_cells[2].text = f"Karakter '{char}' muncul {count} kali dalam teks"

    for i, step in enumerate(steps, 1):
        doc.add_heading(f'Langkah {i}', level=1)
        left, right, merged = step
        doc.add_paragraph(f"Pada langkah ini, kita akan menggabungkan dua node dengan frekuensi terendah:")
        doc.add_paragraph(f"  • Node kiri: {'karakter ' + left.char if left.char else 'gabungan'} dengan frekuensi {left.freq}")
        doc.add_paragraph(f"  • Node kanan: {'karakter ' + right.char if right.char else 'gabungan'} dengan frekuensi {right.freq}")
        doc.add_paragraph(f"Hasil penggabungan adalah node baru dengan frekuensi {merged.freq}.")
        doc.add_paragraph("Berikut adalah visualisasi dari proses penggabungan ini:")
        img_path = create_step_image(step, i)
        doc.add_picture(img_path, width=Inches(3))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Dalam gambar di atas, kita bisa melihat bagaimana dua node digabungkan menjadi satu node baru.")
        doc.add_paragraph("Node di atas adalah hasil penggabungan, sementara dua node di bawahnya adalah node-node yang digabungkan.")
        doc.add_paragraph("Angka '0' pada garis kiri menandakan bahwa untuk mencapai node kiri, kita menggunakan kode '0'.")
        doc.add_paragraph("Sedangkan angka '1' pada garis kanan menandakan bahwa untuk mencapai node kanan, kita menggunakan kode '1'.")
    
    doc.add_heading('Pohon Huffman Akhir', level=1)
    doc.add_paragraph("Setelah semua langkah penggabungan selesai, kita mendapatkan pohon Huffman akhir seperti berikut:")
    final_tree_img = create_final_tree_image(root)
    doc.add_picture(final_tree_img, width=Inches(4))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Dalam pohon Huffman akhir ini, setiap daun (node paling bawah) merepresentasikan sebuah karakter.")
    doc.add_paragraph("Untuk mendapatkan kode Huffman suatu karakter, kita menelusuri jalur dari akar (node paling atas) ke daun karakter tersebut.")
    doc.add_paragraph("Setiap kali kita bergerak ke kiri, kita menambahkan '0' pada kode, dan setiap kali bergerak ke kanan, kita menambahkan '1'.")
    
    doc.add_heading('Kode Huffman', level=1)
    doc.add_paragraph("Berdasarkan pohon Huffman yang telah kita buat, berikut adalah kode Huffman untuk setiap karakter:")
    code_table = doc.add_table(rows=1, cols=3)
    code_table.style = 'Table Grid'
    hdr_cells = code_table.rows[0].cells
    hdr_cells[0].text = 'Karakter'
    hdr_cells[1].text = 'Kode Huffman'
    hdr_cells[2].text = 'Penjelasan'
    for char, code in codes.items():
        row_cells = code_table.add_row().cells
        row_cells[0].text = f"'{char}'"
        row_cells[1].text = code
        row_cells[2].text = f"Untuk mencapai karakter '{char}', kita mengikuti jalur {' kemudian '.join(code)}"
    
    compression_ratio, space_saving = calculate_compression(text, codes)
    doc.add_heading('Analisis Kompresi', level=1)
    doc.add_paragraph(f"Rasio Kompresi: {compression_ratio:.2%}")
    doc.add_paragraph(f"Ini berarti ukuran data setelah kompresi adalah {compression_ratio:.2%} dari ukuran aslinya.")
    doc.add_paragraph(f"Penghematan Ruang: {space_saving:.2%}")
    doc.add_paragraph(f"Kita berhasil menghemat {space_saving:.2%} dari ruang penyimpanan yang dibutuhkan.")
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    doc.save('proses_kode_huffman.docx')

text = "Farhan Aulianda"
root, steps = create_huffman_tree(text)
codes = generate_codes(root)
create_word_document(text, steps, root, codes)
